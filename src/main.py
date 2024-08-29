import os  
import time  
import logging 
import fitz 
from docx import Document as DocxDocument 
from pydantic import BaseModel 
import ast  
import json

from llama_index.llms.ollama import Ollama  
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, Document, PromptTemplate 
from llama_index.core.embeddings import resolve_embed_model  
from llama_index.core.tools import QueryEngineTool, ToolMetadata 
from llama_index.core.agent import ReActAgent
from llama_index.core.query_pipeline import QueryPipeline 
from llama_index.core.output_parsers import PydanticOutputParser 

from prompts import context, code_parser_template

# Configure logging to capture important events and errors
logging.basicConfig(level=logging.INFO)

# Define the Pydantic output model
class CodeOutput(BaseModel):
    code: str
    description: str
    filename: str

def parse_pdf(file_path):
    """Parse PDF files using PyMuPDF."""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def parse_docx(file_path):
    """Parse DOCX files using python-docx."""
    doc = DocxDocument(file_path)
    text = "\n".join(para.text for para in doc.paragraphs)
    return text

class CustomSimpleDirectoryReader(SimpleDirectoryReader):
    """Custom document loader to handle both PDF and DOCX files."""
    def __init__(self, input_dir, file_extractor=None):
        super().__init__(input_dir)
        self.file_extractor = file_extractor or {}

    def load_data(self):
        documents = []
        for root, _, files in os.walk(self.input_dir):
            for file in files:
                file_path = os.path.join(root, file)
                try:
                    if file.endswith('.pdf'):
                        content = parse_pdf(file_path)
                    elif file.endswith('.docx'):
                        content = parse_docx(file_path)
                    else:
                        logging.warning(f"Unsupported file type: {file}")
                        continue

                    document = Document(text=content, metadata={"source": file_path})
                    documents.append(document)
                except Exception as e:
                    logging.error(f"Failed to process file {file_path}: {e}")
                    continue
        return documents

def add_web_resource(documents):
    """Add an additional resource from the Komputronik Biznes website."""
    web_content = "Komputronik Biznes is an IT solutions integrator offering comprehensive services in consulting, hardware, software, and IT infrastructure modernization."
    documents.append(Document(text=web_content, metadata={"source": "https://www.komputronikbiznes.pl/"}))

def setup_llm_model():
    """Setup LLM model with a specific configuration."""
    return Ollama(model="llama3.1", request_timeout=30.0)

def setup_vector_index(documents, embed_model):
    """Create a vector index from the loaded documents."""
    return VectorStoreIndex.from_documents(documents, embed_model=embed_model)

def setup_query_engine(vector_index, llm_model):
    """Convert the vector index into a query engine that can interface with the LLM."""
    return vector_index.as_query_engine(llm=llm_model)

def setup_tools(query_engine):
    """Setup tools to be used by the ReActAgent."""
    return [
        QueryEngineTool(
            query_engine=query_engine,
            metadata=ToolMetadata(
                name="api_documentation",
                description="This provides documentation about a project called DEGA. Use this for reading documentation."
            ),
        )
    ]

def setup_react_agent(tools, llm_model, context):
    """Initialize the ReActAgent using the tools."""
    return ReActAgent.from_tools(tools, llm=llm_model, verbose=True, context=context)

def setup_output_pipeline(json_prompt_tmpl, llm_model):
    """Setup a query pipeline with a chain of templates and models."""
    return QueryPipeline(chain=[json_prompt_tmpl, llm_model])

def handle_query(agent, output_pipeline, prompt, is_code_related):
    """Handle the query, iterate through thought processes, and return the final output."""
    retries = 0
    clean_json = None

    for i in range(3):  # Iterate 3 times to simulate a thought process
        try:
            logging.info(f"Iteration {i+1}/3: Querying the agent.")
            
            if is_code_related:
                result = agent.query(prompt)
                next_result = output_pipeline.run(response=result)
            else:
                result = agent.query(prompt)
                next_result = result

            print(f"Raw output (Iteration {i+1}/3): {next_result}")

            if is_code_related:
                json_str = str(next_result).replace("assistant:", "").strip()
                clean_json = json.loads(json_str)
            else:
                clean_json = {
                    "code": "",
                    "description": str(next_result),
                    "filename": "response.txt"
                }

            logging.info(f"Iteration {i+1}/3 complete.")
        except (json.JSONDecodeError, ValueError, SyntaxError) as e:
            retries += 1
            logging.error(f"Parsing error: {e}. Retrying... ({retries}/3)")
        except Exception as e:
            retries += 1
            logging.error(f"Unexpected error: {e}. Retrying... ({retries}/3)")

        if clean_json:
            break

    if retries == 3:
        logging.error("Failed to get a response after 3 retries. Please try again.")

    return clean_json

def save_output_to_file(clean_json):
    """Save the output to a file."""
    filename = clean_json["filename"]

    try:
        with open(filename, "w") as f:
            f.write(clean_json["code"] if clean_json["code"] else clean_json["description"])
        logging.info(f"Output saved to {filename}")
    except Exception as e:
        logging.error(f"Error saving output to file: {e}")

def main():
    llm_model = setup_llm_model()

    # Load documents from a directory
    documents = CustomSimpleDirectoryReader("/home/nikodem-ub1/github/RAG-client/data").load_data()
    add_web_resource(documents)

    embed_model = resolve_embed_model("local:BAAI/bge-m3")
    vector_index = setup_vector_index(documents, embed_model)
    query_engine = setup_query_engine(vector_index, llm_model)

    tools = setup_tools(query_engine)
    agent = setup_react_agent(tools, llm_model, context)

    parser = PydanticOutputParser(CodeOutput)
    json_prompt_str = parser.format(code_parser_template)
    json_prompt_tmpl = PromptTemplate(json_prompt_str)

    output_pipeline = setup_output_pipeline(json_prompt_tmpl, llm_model)

    while (prompt := input("Enter a prompt (q to quit) -> ")) != "q":
        start_time = time.time()

        # Check if the prompt is likely asking for code generation
        is_code_related = any(keyword in prompt.lower() for keyword in ["code", "script", "function", "class", "method"])

        clean_json = handle_query(agent, output_pipeline, prompt, is_code_related)

        if clean_json:
            print("Final Output:")
            print(clean_json["code"] if clean_json["code"] else clean_json["description"])
            print("\n\nDescription:", clean_json["description"])

            save_output_to_file(clean_json)

        end_time = time.time()
        logging.info(f"Time taken: {end_time - start_time} seconds")

# Entry point of the script
if __name__ == "__main__":
    main()
