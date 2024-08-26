import os  # Provides functions for interacting with the operating system
import time  # Used for time-related functions such as tracking execution time
import logging  # Allows for configurable logging of messages for debugging and monitoring
import fitz  # PyMuPDF for PDF parsing, allows extraction of text from PDF files
from docx import Document as DocxDocument  # python-docx for DOCX parsing, allows extraction of text from DOCX files
from pydantic import BaseModel  # Pydantic for data validation and settings management using Python type annotations
import ast  # Provides functions to safely evaluate strings containing Python literals

from llama_index.llms.ollama import Ollama  # Import Ollama LLM interface for interacting with LLM models
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, Document, PromptTemplate  # Core components for vector indexing, reading documents, and templating prompts
from llama_index.core.embeddings import resolve_embed_model  # Resolves and loads embedding models for converting text to vectors
from llama_index.core.tools import QueryEngineTool, ToolMetadata  # Provides tools and metadata for querying LLMs
from llama_index.core.agent import ReActAgent  # ReActAgent is used for creating agents that can perform complex tasks using LLMs
from llama_index.core.query_pipeline import QueryPipeline  # Allows chaining of multiple LLM operations into a query pipeline
from llama_index.core.output_parsers import PydanticOutputParser  # Parser for processing LLM outputs into structured data using Pydantic models

from prompts import context, code_parser_template

# Configure logging to capture important events and errors
logging.basicConfig(level=logging.INFO)

# Function to parse PDF files using PyMuPDF
# This function reads and extracts text from each page of the PDF file
def parse_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()  # Extract text from each page and append to a string
    return text

# Function to parse DOCX files using python-docx
# This function reads and extracts text from each paragraph of the DOCX file
def parse_docx(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join(para.text for para in doc.paragraphs)  # Join paragraphs with newline
    return text

# Custom document loader to handle both PDF and DOCX files
# This class extends SimpleDirectoryReader to include custom file parsing logic
class CustomSimpleDirectoryReader(SimpleDirectoryReader):
    def __init__(self, input_dir, file_extractor=None):
        super().__init__(input_dir)  # Initialize the parent class with the input directory
        self.file_extractor = file_extractor or {}  # Dictionary to hold custom file extractors

    # Overriding the load_data method to support PDF and DOCX files
    def load_data(self):
        documents = []
        # Walk through the directory to find and process files
        for root, _, files in os.walk(self.input_dir):
            for file in files:
                file_path = os.path.join(root, file)  # Full path to the file
                try:
                    if file.endswith('.pdf'):
                        content = parse_pdf(file_path)  # Parse PDF files
                    elif file.endswith('.docx'):
                        content = parse_docx(file_path)  # Parse DOCX files
                    else:
                        logging.warning(f"Unsupported file type: {file}")
                        continue  # Skip unsupported file types

                    # Create a Document object with the extracted content and metadata
                    document = Document(text=content, metadata={"source": file_path})
                    documents.append(document)
                except Exception as e:
                    logging.error(f"Failed to process file {file_path}: {e}")
                    continue  # Skip files that cause errors
        return documents

# Adding an additional resource from the Komputronik Biznes website
def add_web_resource():
    web_content = "Komputronik Biznes is an IT solutions integrator offering comprehensive services in consulting, hardware, software, and IT infrastructure modernization."
    documents.append(Document(text=web_content, metadata={"source": "https://www.komputronikbiznes.pl/"}))

# Setup LLM model with a specific configuration
# Ollama is used to initialize the LLM model with a request timeout
model = Ollama(model="llama3.1", request_timeout=30.0)

# Load documents from a directory, using the custom document loader
documents = CustomSimpleDirectoryReader("/home/nikodem-ub1/github/RAG-client/data").load_data()

# Add web resource content
add_web_resource()

# Resolve the embedding model using the specified model name
# Embedding models are used to convert text data into numerical vectors
embed_model = resolve_embed_model("local:BAAI/bge-m3")

# Create a vector index from the loaded documents
# VectorStoreIndex is used for efficient text search and retrieval
vector_index = VectorStoreIndex.from_documents(documents, embed_model=embed_model)

# Convert the vector index into a query engine that can interface with the LLM
query_engine = vector_index.as_query_engine(llm=model)

# Setup tools
# These tools will be used by the agent to interact with the LLM and perform specific tasks
tools = [
    QueryEngineTool(
        query_engine=query_engine,
        metadata=ToolMetadata(
            name="api_documentation",
            description="This provides documentation about a project called DEGA. Use this for reading documentation."
        ),
    )
]

# Initialize the ReActAgent using the tools
agent = ReActAgent.from_tools(tools, llm=model, verbose=True, context=context)

# Pydantic output formatter
# This section defines the structure of the expected output using Pydantic
class CodeOutput(BaseModel):
    code: str  # The generated code as a string
    description: str  # Description of the code
    filename: str  # Suggested filename for saving the code

# Initialize the Pydantic output parser with the defined output structure
parser = PydanticOutputParser(CodeOutput)

# Format the code parser template with the parser
json_prompt_str = parser.format(code_parser_template)

# Create a PromptTemplate object using the formatted string
json_prompt_tmpl = PromptTemplate(json_prompt_str)

# Setup a query pipeline with a chain of templates and models
# The pipeline processes the model's output through multiple stages
output_pipeline = QueryPipeline(chain=[json_prompt_tmpl, model])

# Main loop to interact with the user
def main():
    while (prompt := input("Enter a prompt (q to quit) -> ")) != "q":
        start_time = time.time()  # Record the start time for performance measurement

        retries = 0  # Initialize retry counter
        while retries < 3:  # Allow up to 3 retries in case of errors
            try:
                result = agent.query(prompt)  # Query the LLM agent with the user prompt
                next_result = output_pipeline.run(response=result)  # Process the agent's response through the pipeline

                # Safely parse JSON output
                clean_json = ast.literal_eval(str(next_result).replace("assistant:", ""))
                break  # Exit the retry loop if successful
            except (ValueError, SyntaxError) as e:
                retries += 1
                logging.error(f"Parsing error: {e}. Retrying... ({retries}/3)")
            except Exception as e:
                retries += 1
                logging.error(f"Unexpected error: {e}. Retrying... ({retries}/3)")

        if retries == 3:  # If all retries failed, inform the user
            logging.error("Failed to get a response after 3 retries. Please try again.")
            continue  # Skip to the next iteration of the loop

        # Display the generated code and its description
        print("Code Output:")
        print(clean_json["code"])
        print("\n\nDescription:", clean_json["description"])

        filename = clean_json["filename"]  # Extract the suggested filename

        try:
            # Attempt to save the generated code to a file
            with open(filename, "w") as f:
                f.write(clean_json["code"])
            logging.info(f"Code saved to {filename}")
        except Exception as e:
            logging.error(f"Error saving code to file: {e}")

        end_time = time.time()  # Record the end time for performance measurement
        logging.info(f"Time taken: {end_time - start_time} seconds")  # Log the time taken

# Entry point of the script
if __name__ == "__main__":
    main()
